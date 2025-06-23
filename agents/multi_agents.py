import os
import sys
import re
import io
import asyncio
import PyPDF2
import docx
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from langgraph_supervisor import create_supervisor
from agents.filesystem_agent import FilesystemAgent
from agents.text_extraction_agent import TextExtractionAgent
from agents.file_classification_agent import FileClassificationAgent
from agents.metadata_agent import MetadataAgent
from config.llm import gemini
from utils.pretty_print_message import pretty_print_messages
from PIL import Image


async def create_supervisor_agent():
    filesystem_agent = await FilesystemAgent.create()
    text_extraction_agent = TextExtractionAgent()
    file_classification_agent = FileClassificationAgent()
    metadata_agent = MetadataAgent()
    supervisor = create_supervisor(
        agents=[filesystem_agent, text_extraction_agent, file_classification_agent, metadata_agent],
        model= gemini,
        prompt = """
        Bạn đóng vai trò là người điều phối trong một hệ thống gồm nhiều Agent chuyên trách xử lý tài liệu.  
        Nhiệm vụ của bạn là tiếp nhận và phân tích yêu cầu đầu vào (có thể là truy vấn, tệp tin hoặc mô tả nhiệm vụ), sau đó phân công chính xác các Agent phù hợp để thực hiện tác vụ đó.

        Khi tiếp nhận đầu vào, bạn cần thực hiện theo quy trình sau (trừ khi được yêu cầu cụ thể khác):

        1. Nếu đầu vào là truy vấn tìm tài liệu: dùng Filesystem Agent để truy xuất danh sách file phù hợp.
        2. Nếu đã có danh sách file: sử dụng Text Extraction Agent để trích xuất nội dung văn bản.
        3. Sau khi có nội dung: sử dụng File Classification Agent để phân loại tài liệu.
        4. Nếu tài liệu đã được phân loại: sử dụng Metadata Agent để trích xuất thông tin và tổng hợp metadata.

        Lưu ý:
        - Không gọi Metadata Agent nếu chưa phân loại tài liệu.
        - Chỉ gọi Text Extraction nếu đã có file.
        - Chỉ gọi File Classification nếu đã có nội dung.
        """,
        supervisor_name="SupervisorAgent"
    ).compile()
    return supervisor


async def create_graph():
    supervisor = await create_supervisor_agent()
    png_bytes = supervisor.get_graph().draw_mermaid_png()  # giả sử trả về bytes PNG

    # Lưu bytes PNG ra file
    with open("graph.png", "wb") as f:
        f.write(png_bytes)

    # Mở và hiển thị ảnh bằng PIL
    img = Image.open(io.BytesIO(png_bytes))
    img.show()

async def process_direct_workflow(query):
    """Process a workflow by directly calling each agent in sequence"""
    print("\n=== Starting Direct Workflow ===\n")
    print(f"Query: {query}\n")
    
    # Step 1: Use FilesystemAgent to find files
    print("Step 1: Finding files with FilesystemAgent...")
    filesystem_agent = await FilesystemAgent.create()
    files_result = await filesystem_agent.run(query, "direct_workflow")
    
    if isinstance(files_result, dict) and 'content' in files_result:
        files_content = files_result['content']
    else:
        files_content = str(files_result)
        
    print(f"\nFiles found:\n{files_content}\n")
    
    # Step 2: Extract text from files
    print("\nStep 2: Extracting text with TextExtractionAgent...")
    text_extraction_agent = TextExtractionAgent()
    session_id = "direct_workflow"
    
    # Extract file paths from the result
    import re
    file_paths = re.findall(r'[A-Z]:\\\\?(?:[^\\/:*?"<>|\r\n]+\\\\?)*[^\\/:*?"<>|\r\n]*\.[a-zA-Z0-9]+', files_content)
    
    # If regex didn't work, try a simpler approach
    if not file_paths:
        print("Using alternative file path extraction method...")
        # Look for common file extensions
        for ext in ['.pdf', '.docx', '.pptx', '.xlsx', '.txt']:
            if ext in files_content:
                start_idx = files_content.find('C:')
                if start_idx != -1:
                    end_idx = files_content.find('\n', start_idx)
                    if end_idx == -1:  # If no newline, take the rest of the string
                        end_idx = len(files_content)
                    file_path = files_content[start_idx:end_idx].strip()
                    file_paths = [file_path]
                    break
    
    if not file_paths:
        print("No file paths found in the FilesystemAgent result.")
        print(f"Raw result: {files_content}")
        return None  # Return None explicitly
        
    # Extract text from the first file
    file_path = file_paths[0]
    print(f"Extracting text from: {file_path}")
    
    # Use TextExtractionAgent to extract text
    text_extraction_agent = TextExtractionAgent()
    
    print(f"Using TextExtractionAgent to extract text from: {file_path}")
    extraction_query = f"Extract text from {file_path}"
    try:
        text_content = text_extraction_agent.invoke(extraction_query, session_id)
        if not text_content or len(text_content.strip()) < 10:
            print("TextExtractionAgent returned empty or very short content. Trying direct extraction...")
            # Fallback to direct extraction
            if file_path.lower().endswith('.pdf'):
                import PyPDF2
                try:
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        text_content = ""
                        for page_num in range(len(reader.pages)):
                            text_content += reader.pages[page_num].extract_text() + "\n"
                    print("Successfully extracted text using PyPDF2")
                except Exception as e:
                    print(f"Error extracting PDF with PyPDF2: {e}")
            elif file_path.lower().endswith('.docx'):
                import docx
                try:
                    doc = docx.Document(file_path)
                    text_content = ""
                    for para in doc.paragraphs:
                        text_content += para.text + "\n"
                    print("Successfully extracted text using python-docx")
                except Exception as e:
                    print(f"Error extracting DOCX: {e}")
    except Exception as e:
        print(f"Error using TextExtractionAgent: {e}")
        # Fallback to direct extraction
        if file_path.lower().endswith('.pdf'):
            import PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text_content = ""
                    for page_num in range(len(reader.pages)):
                        text_content += reader.pages[page_num].extract_text() + "\n"
                print("Successfully extracted text using PyPDF2")
            except Exception as e:
                print(f"Error extracting PDF with PyPDF2: {e}")
                text_content = "Error extracting text from PDF"
        elif file_path.lower().endswith('.docx'):
            import docx
            try:
                doc = docx.Document(file_path)
                text_content = ""
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
                print("Successfully extracted text using python-docx")
            except Exception as e:
                print(f"Error extracting DOCX: {e}")
                text_content = "Error extracting text from DOCX"
        else:
            text_content = f"Could not extract text from {file_path}"
    
    print(f"\nExtracted text (first 500 chars):\n{text_content[:500]}...\n")
    
    # Step 3: Classify the document
    print("\nStep 3: Classifying document with FileClassificationAgent...")
    file_classification_agent = FileClassificationAgent()
    classification_query = f"Classify the following document:\n{text_content[:2000]}"
    classification_result = file_classification_agent.invoke(classification_query, session_id)
    
    print(f"\nClassification result:\n{classification_result}\n")
    
    # Step 4: Generate metadata
    print("\nStep 4: Generating metadata with MetadataAgent...")
    metadata_agent = MetadataAgent()
    metadata_query = f"Generate metadata for the following document:\nFile: {file_path}\nClassification: {classification_result}\nContent: {text_content[:1000]}"
    metadata_result = metadata_agent.invoke(metadata_query, session_id)
    
    print(f"\nMetadata result:\n{metadata_result}\n")
    print("\n=== Workflow Complete ===\n")
    
    return {
        "files": files_content,
        "text": text_content[:500] + "...",
        "classification": classification_result,
        "metadata": metadata_result
    }

async def main():
    # Choose which approach to use
    use_direct_workflow = True
    
    if use_direct_workflow:
        # Direct workflow approach
        result = await process_direct_workflow("Tìm kiếm các file liên quan tới project-final trong thư mục cho phép")
        
        if result is not None:
            print("\nWorkflow Summary:")
            for key, value in result.items():
                print(f"\n{key.upper()}:\n{value[:200]}..." if isinstance(value, str) and len(value) > 200 else f"\n{key.upper()}:\n{value}")
        else:
            print("\nWorkflow failed to complete. No results to display.")
            print("Try running with a different query or check agent configurations.")
            # Try a more specific query
            print("\nTrying with a more specific query...")
            result = await process_direct_workflow("Find the file project-final-DV.pdf in the data directory")
            
            if result is not None:
                print("\nWorkflow Summary:")
                for key, value in result.items():
                    print(f"\n{key.upper()}:\n{value[:200]}..." if isinstance(value, str) and len(value) > 200 else f"\n{key.upper()}:\n{value}")
            else:
                print("\nSecond attempt also failed. Please check agent configurations.")
                # Fall back to supervisor approach
                print("\nFalling back to supervisor approach...")
                supervisor = await create_supervisor_agent()
                async for chunk in supervisor.astream(
                    {
                        "messages": [
                            {
                                "role": "user",
                                "content": "Tìm kiếm các file liên quan tới project-final trong thư mục cho phép và xuất metadata vào file excel."
                            }
                        ]
                    }
                ):
                    pretty_print_messages(chunk)
    else:
        # Supervisor approach
        supervisor = await create_supervisor_agent()
        async for chunk in supervisor.astream(
            {
                "messages": [
                    {
                        "role": "user",
                        "content": "Tìm kiếm các file liên quan tới project-final trong thư mục cho phép và xuất metadata vào file excel."
                    }
                ]
            }
        ):
            pretty_print_messages(chunk)

if __name__ == "__main__":
    asyncio.run(main())